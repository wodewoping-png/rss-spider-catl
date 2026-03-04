# -*- coding: utf-8 -*-
"""
classify_weekly_onefile.py

Single-file weekly classifier for GitHub Actions.

What it does:
1) Pick latest weekly CSV from output/weekly (or -i path).
2) Build stable per-record ID (sha1 of link / fallback fields).
3) Build text_for_classify = title + abstract(head+tail).
4) 3-stage labeling:
   A) Hard exclusion overrides (metal-air/flow/etc -> 其他储能器件; methanol/ammonia/hydrogen synthesis -> 氢氨醇 overrides CCUS)
   B) Keyword strong routing (from classification.txt "关键词：" lists) + simple scoring
   C) Transformer semantic routing (sentence-transformers embeddings) -> strong auto-label or weak candidates
5) GPT by_id classification (only for unresolved items). Output mapping by_id, retry missing once.
6) Postprocess exclusions again + choose primary label + review flags.
7) Write XLSX with:
   - one sheet per category
   - 全部 / 未分类 / 待复核 / 仍缺失标签
8) Write DOCX grouped by category (simple bilingual-ish: ZH fields if present; otherwise EN).

Required files:
- classification.txt (each line "类名：说明...关键词：k1, k2, ..."; you said you already updated and added “其他”)
- requirements.txt should include:
  pandas, openpyxl, python-docx, openai, sentence-transformers, scikit-learn, torch

Env:
- OPENAI_API_KEY (required if GPT stage is used)
- OPENAI_MODEL (default gpt-4o-mini)

Tips:
- If you only want to test non-GPT routing: pass --skip-gpt
- If you want GPT for all items: pass --force-gpt
"""

import argparse
import hashlib
import json
import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# -----------------------
# Config (ENV + defaults)
# -----------------------
OPENAI_MODEL = (os.getenv("OPENAI_MODEL") or "").strip() or "gpt-4o-mini"
OPENAI_TEMPERATURE = float(os.getenv("OPENAI_TEMPERATURE", "0"))
OPENAI_MAX_RETRIES = int(os.getenv("OPENAI_MAX_RETRIES", "3"))
OPENAI_BASE_URL = (os.getenv("OPENAI_BASE_URL") or "").strip()

CLASSIFY_BATCH_SIZE = int(os.getenv("CLASSIFY_BATCH_SIZE", "12"))
GPT_RETRY_MISSING_ONCE = True

TEXT_HEAD_CHARS = int(os.getenv("TEXT_HEAD_CHARS", "800"))
TEXT_TAIL_CHARS = int(os.getenv("TEXT_TAIL_CHARS", "500"))
TEXT_MAX_CHARS = int(os.getenv("TEXT_MAX_CHARS", "1600"))

# Keyword routing thresholds
KEYWORD_MIN_HITS_STRONG = int(os.getenv("KEYWORD_MIN_HITS_STRONG", "2"))  # >=2 hits -> strong route
KEYWORD_MIN_HITS_WEAK = int(os.getenv("KEYWORD_MIN_HITS_WEAK", "1"))      # 1 hit -> weak (candidate)

# Embedding routing thresholds (cosine similarity)
EMB_STRONG_THRESHOLD = float(os.getenv("EMB_STRONG_THRESHOLD", "0.55"))
EMB_WEAK_THRESHOLD = float(os.getenv("EMB_WEAK_THRESHOLD", "0.45"))
EMB_TOPK_CANDIDATES = int(os.getenv("EMB_TOPK_CANDIDATES", "3"))

# Review threshold (only meaningful if GPT returns confidence; we also derive confidence from routing)
REVIEW_CONFIDENCE_THRESHOLD = float(os.getenv("REVIEW_CONFIDENCE_THRESHOLD", "0.55"))

DEFAULT_CLASSIFICATION_FILE = os.getenv("CLASSIFICATION_FILE", "classification.txt")

DEFAULT_INPUT_DIR = os.getenv("WEEKLY_DIR", "output/weekly")
DEFAULT_DEBUG_DIR = os.getenv("DEBUG_DIR", "output/debug")

# Sheets
SHEET_ALL = "全部"
SHEET_UNLABELED = "未分类"
SHEET_REVIEW = "待复核"
SHEET_STILL_MISSING = "仍缺失标签"

print("[boot] classify_weekly_onefile.py started", flush=True)


# -----------------------
# Helpers
# -----------------------
def _normalize_cell(v) -> str:
    if v is None:
        return ""
    return str(v).strip()


def _fmt_secs(s: float) -> str:
    s = max(0, int(s))
    h = s // 3600
    m = (s % 3600) // 60
    sec = s % 60
    if h > 0:
        return f"{h}h {m}m {sec}s"
    if m > 0:
        return f"{m}m {sec}s"
    return f"{sec}s"


def _chunked(items, n):
    for i in range(0, len(items), n):
        yield items[i:i + n]


def _clean_json_text(s: str) -> str:
    s = (s or "").strip()
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.IGNORECASE)
        s = re.sub(r"\s*```$", "", s)
    return s.strip()


def _safe_sheet_name(name: str, used: set) -> str:
    base = re.sub(r"[\\/*?:\[\]]", "_", _normalize_cell(name)) or "Sheet"
    base = base[:31]
    candidate = base
    i = 2
    while candidate in used:
        suffix = f"_{i}"
        candidate = base[: 31 - len(suffix)] + suffix
        i += 1
    used.add(candidate)
    return candidate


def _dump_debug(tag: str, payload: dict, raw_text: str = "", reason: str = ""):
    outdir = Path(DEFAULT_DEBUG_DIR)
    outdir.mkdir(parents=True, exist_ok=True)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    stem = f"{tag}_{ts}_{hashlib.sha1((reason+raw_text).encode('utf-8', errors='ignore')).hexdigest()[:8]}"
    (outdir / f"{stem}_reason.txt").write_text(reason or "", encoding="utf-8")
    (outdir / f"{stem}_payload.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    if raw_text:
        (outdir / f"{stem}_raw.txt").write_text(raw_text, encoding="utf-8")
    print(f"[debug] dumped to {outdir.resolve()}", flush=True)


# -----------------------
# Input discovery
# -----------------------
def _extract_date_from_name(filename: str):
    stem = Path(filename).stem
    m = re.search(r"(\d{4})[-_]?(\d{2})[-_]?(\d{2})", stem)
    if not m:
        return None
    y, mo, d = map(int, m.groups())
    try:
        return datetime(y, mo, d).date()
    except ValueError:
        return None


def pick_latest_weekly_csv(folder: str) -> Path:
    folder = Path(folder)
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder.resolve()}")

    prefixes = ["weekly_news_with_abstract_", "news_with_abstract_"]
    candidates = []
    for p in folder.glob("*.csv"):
        if "_translated" in p.stem:
            continue
        if any(p.name.startswith(pref) for pref in prefixes):
            candidates.append(p)

    if not candidates:
        raise FileNotFoundError(f"No weekly CSV found in {folder.resolve()}")

    with_dates, without_dates = [], []
    for p in candidates:
        d = _extract_date_from_name(p.name)
        if d is not None:
            with_dates.append((d, p))
        else:
            without_dates.append(p)

    if with_dates:
        with_dates.sort(key=lambda x: (x[0], x[1].stat().st_mtime), reverse=True)
        return with_dates[0][1]
    return max(without_dates, key=lambda p: p.stat().st_mtime)


def ensure_base_columns(df: pd.DataFrame) -> pd.DataFrame:
    needed = [
        "title",
        "link",
        "published",
        "source",
        "pub_date",
        "doi",
        "abstract",
        "abstract_source",
        "must_have_abstract",
        "title_zh",
        "abstract_zh",
        "categories",
        "id",
        "primary_category",
        "confidence",
        "route",
        "need_review",
    ]
    for c in needed:
        if c not in df.columns:
            df[c] = ""
        df[c] = df[c].fillna("").astype(str)
    return df


def make_stable_id(row: dict) -> str:
    link = _normalize_cell(row.get("link", ""))
    if link:
        base = link
    else:
        base = "|".join([
            _normalize_cell(row.get("title", "")),
            _normalize_cell(row.get("source", "")),
            _normalize_cell(row.get("pub_date", "")) or _normalize_cell(row.get("published", "")),
        ])
    base = base.strip() or json.dumps(row, ensure_ascii=False)
    return hashlib.sha1(base.encode("utf-8", errors="ignore")).hexdigest()[:12]


# -----------------------
# Text prep
# -----------------------
def strip_leading_abstract(text: str) -> str:
    if not text:
        return text
    s = text.lstrip()
    return re.sub(r"^(abstract)\s*[:.\-–—]*\s*", "", s, flags=re.IGNORECASE)


def build_text_for_classify(title: str, abstract: str) -> str:
    title = _normalize_cell(title)
    abstract = _normalize_cell(abstract)
    abstract = strip_leading_abstract(abstract)

    if not abstract:
        text = title
    else:
        a = abstract.replace("\r\n", "\n").replace("\r", "\n")
        a = re.sub(r"\s+", " ", a).strip()
        head = a[:TEXT_HEAD_CHARS]
        tail = a[-TEXT_TAIL_CHARS:] if len(a) > TEXT_HEAD_CHARS else ""
        parts = [title] if title else []
        if head:
            parts.append(head)
        if tail and tail != head:
            parts.append("... " + tail)
        text = "\n".join([p for p in parts if p])

    text = text.strip()
    if len(text) > TEXT_MAX_CHARS:
        text = text[:TEXT_MAX_CHARS]
    return text


# -----------------------
# classification.txt parsing (names + desc + keywords)
# -----------------------
def parse_classification_file(path: str) -> Tuple[List[str], Dict[str, str], Dict[str, List[str]]]:
    """
    Returns:
      ordered_categories: [cat1, cat2, ...]
      cat_desc: {cat: desc_text}
      cat_keywords: {cat: [kw1, kw2, ...]} (parsed from "关键词：")
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"classification file not found: {p.resolve()}")

    ordered = []
    desc_map = {}
    kw_map: Dict[str, List[str]] = {}

    for raw in p.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue

        # split "类名：..."
        if "：" in line:
            name, rest = line.split("：", 1)
        elif ":" in line:
            name, rest = line.split(":", 1)
        else:
            name, rest = line, ""

        cat = name.strip()
        if not cat:
            continue

        # Extract keywords (关键词：...)
        rest_clean = rest.strip()
        kws = []
        m = re.search(r"(关键词|keywords)\s*[:：]\s*(.+)$", rest_clean, flags=re.IGNORECASE)
        if m:
            kw_part = m.group(2).strip()
            # Split by commas / Chinese commas / semicolons
            kws = [x.strip() for x in re.split(r"[，,;；\n]+", kw_part) if x.strip()]
            # Remove trailing punctuation
            kws = [re.sub(r"[。\.]+$", "", x) for x in kws]

            # Description is text before keywords marker
            desc_text = rest_clean[:m.start()].strip()
        else:
            desc_text = rest_clean

        ordered.append(cat)
        desc_map[cat] = desc_text
        kw_map[cat] = kws

    if not ordered:
        raise RuntimeError("No valid categories parsed from classification.txt")

    return ordered, desc_map, kw_map


# -----------------------
# Deterministic exclusion / override rules (hard-coded)
# -----------------------
# Strong signals for "其他储能器件" (exclude Li-ion components)
PAT_OTHER_STORAGE = [
    r"\bmetal[- ]air\b",
    r"\bli[- ]air\b",
    r"\bzn[- ]air\b",
    r"\b(al|zn|li|na|k|mg|ca)[- ]air\b",
    r"\bmetal[- ]oxygen\b",
    r"\boxygen battery\b",
    r"\bmetal[- ]co2\b",
    r"\bco2 battery\b",
    r"\bflow battery\b",
    r"\bredox flow\b",
    r"\bvanadium flow\b",
    r"\blead[- ]acid\b",
    r"\b(supercapacitor|ultracapacitor)\b",
]

# Strong signals for "氢氨醇" that should override CCUS
PAT_HAM_STRONG = [
    r"\bwater electrolysis\b",
    r"\belectrolyser\b",
    r"\bgreen hydrogen\b",
    r"\bhydrogen production\b",
    r"\bammonia synthesis\b",
    r"\bhaber[- ]bosch\b",
    r"\bmethanol synthesis\b",
    r"\bco2 hydrogenation\b.*\bmethanol\b",
    r"\bpower[- ]to[- ]x\b",
    r"\bpower[- ]to[- ]ammonia\b",
    r"\bpower[- ]to[- ]methanol\b",
]

# Strong signals for CCUS capture/storage (non-HAM)
PAT_CCUS_STRONG = [
    r"\bcarbon capture\b",
    r"\bco2 capture\b",
    r"\bdirect air capture\b",
    r"\bdac\b",
    r"\bco2 storage\b",
    r"\bgeologic(al)? storage\b",
    r"\bamine sorbent\b",
    r"\bcarbon mineralization\b",
    r"\bco2 mineralization\b",
]


def _regex_any(patterns: List[str], text: str) -> bool:
    for pat in patterns:
        if re.search(pat, text, flags=re.IGNORECASE):
            return True
    return False


def apply_hard_overrides(text: str, labels: List[str], ordered_categories: List[str]) -> Tuple[List[str], Optional[str]]:
    """
    Apply deterministic override rules.
    Returns (new_labels, reason_if_changed)
    """
    t = text or ""
    changed_reason = None
    new = list(labels or [])

    # Force other storage if strong match
    if _regex_any(PAT_OTHER_STORAGE, t):
        # Keep only "其他储能器件" if it exists in your classification
        if "其他储能器件" in ordered_categories:
            new = ["其他储能器件"]
            changed_reason = "hard_override:other_storage"
            return new, changed_reason

    # HAM overrides CCUS if HAM strong
    if _regex_any(PAT_HAM_STRONG, t):
        if "氢氨醇" in ordered_categories:
            if "CCUS" in new:
                new = [x for x in new if x != "CCUS"]
            if "氢氨醇" not in new:
                new.append("氢氨醇")
            changed_reason = "hard_override:ham_over_ccus"
            return new, changed_reason

    # If CCUS signals exist and no HAM strong, gently add CCUS if missing
    if _regex_any(PAT_CCUS_STRONG, t) and not _regex_any(PAT_HAM_STRONG, t):
        if "CCUS" in ordered_categories and "CCUS" not in new:
            new.append("CCUS")
            changed_reason = "hard_add:ccus"
            return new, changed_reason

    return new, changed_reason


# -----------------------
# Keyword routing (from classification keywords)
# -----------------------
def keyword_route(text: str, ordered_categories: List[str], cat_keywords: Dict[str, List[str]]) -> Tuple[List[str], float, str]:
    """
    Returns (labels, confidence, route_tag)
    labels empty means no strong decision.
    """
    t = (text or "").lower()
    if not t:
        return [], 0.0, "kw:none"

    # special: if hard other-storage patterns match, skip kw and let hard overrides handle
    if _regex_any(PAT_OTHER_STORAGE, t):
        return [], 0.0, "kw:skip_other_storage"

    scores: Dict[str, int] = {c: 0 for c in ordered_categories}
    hits_detail: Dict[str, List[str]] = {c: [] for c in ordered_categories}

    for cat in ordered_categories:
        kws = cat_keywords.get(cat, []) or []
        if not kws:
            continue
        for kw in kws:
            k = kw.strip()
            if not k:
                continue
            # simple contains match; for very short keywords, require word boundary
            if len(k) <= 3 and re.search(rf"\b{re.escape(k.lower())}\b", t):
                scores[cat] += 1
                hits_detail[cat].append(k)
            else:
                if k.lower() in t:
                    scores[cat] += 1
                    hits_detail[cat].append(k)

    # pick best
    best_cat = max(scores, key=lambda c: scores[c])
    best_score = scores[best_cat]
    if best_score < KEYWORD_MIN_HITS_WEAK:
        return [], 0.0, "kw:none"

    # check tie
    tied = [c for c, s in scores.items() if s == best_score and s > 0]
    if len(tied) > 1:
        # tie -> weak candidates only
        return [], 0.0, f"kw:tie({','.join(tied[:3])})"

    # strong route
    if best_score >= KEYWORD_MIN_HITS_STRONG:
        conf = min(0.92, 0.70 + 0.08 * best_score)
        return [best_cat], conf, f"kw:strong({best_cat}:{best_score})"

    # weak: treat as candidate (not final) -> return empty labels but route info
    return [], 0.0, f"kw:weak_candidate({best_cat}:{best_score})"


def keyword_candidates(text: str, ordered_categories: List[str], cat_keywords: Dict[str, List[str]], topk: int = 3) -> List[str]:
    t = (text or "").lower()
    if not t:
        return []
    scores: Dict[str, int] = {}
    for cat in ordered_categories:
        score = 0
        for kw in (cat_keywords.get(cat, []) or []):
            k = kw.strip()
            if not k:
                continue
            if len(k) <= 3:
                if re.search(rf"\b{re.escape(k.lower())}\b", t):
                    score += 1
            else:
                if k.lower() in t:
                    score += 1
        if score > 0:
            scores[cat] = score
    if not scores:
        return []
    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return [c for c, _ in ranked[:topk]]


# -----------------------
# Transformer semantic routing (SentenceTransformer + cosine sim)
# -----------------------
def try_load_embedder():
    try:
        from sentence_transformers import SentenceTransformer  # type: ignore
        return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
    except Exception as e:
        print(f"[embed] disabled (failed to load sentence-transformers model): {e}", flush=True)
        return None


def cosine_sim_matrix(a, b):
    # a: (n,d), b: (m,d)
    try:
        from sklearn.metrics.pairwise import cosine_similarity  # type: ignore
        return cosine_similarity(a, b)
    except Exception:
        # fallback pure python
        import math
        out = []
        for va in a:
            row = []
            na = math.sqrt(sum(x * x for x in va)) + 1e-12
            for vb in b:
                nb = math.sqrt(sum(x * x for x in vb)) + 1e-12
                dot = sum(x * y for x, y in zip(va, vb))
                row.append(dot / (na * nb))
            out.append(row)
        return out


def build_category_prototypes(ordered_categories: List[str], cat_desc: Dict[str, str], cat_keywords: Dict[str, List[str]]) -> List[str]:
    protos = []
    for c in ordered_categories:
        desc = _normalize_cell(cat_desc.get(c, ""))
        kws = cat_keywords.get(c, []) or []
        # Prototype: "类名：desc 关键词：k1, k2, ..."
        if kws:
            proto = f"{c}：{desc} 关键词：{', '.join(kws[:50])}"
        else:
            proto = f"{c}：{desc}"
        protos.append(proto.strip())
    return protos


def embedding_route_single(
    text: str,
    ordered_categories: List[str],
    cat_emb,
    embedder,
) -> Tuple[List[str], float, str, List[Tuple[str, float]]]:
    """
    Returns (final_labels, confidence, route_tag, top_scores)
    - if strong: final_labels = [top1]
    - else final_labels = []
    """
    if not embedder or cat_emb is None:
        return [], 0.0, "emb:disabled", []

    if not text:
        return [], 0.0, "emb:empty", []

    emb = embedder.encode([text], normalize_embeddings=True)
    sims = cosine_sim_matrix(emb, cat_emb)[0]  # list of floats
    pairs = list(zip(ordered_categories, sims))
    pairs.sort(key=lambda x: x[1], reverse=True)

    top1_cat, top1 = pairs[0]
    tag = f"emb:top1({top1_cat}:{top1:.3f})"

    if top1 >= EMB_STRONG_THRESHOLD:
        # strong auto label
        conf = min(0.90, 0.55 + (top1 - EMB_STRONG_THRESHOLD) * 0.8)
        return [top1_cat], conf, f"emb:strong({top1_cat}:{top1:.3f})", pairs[:5]

    return [], 0.0, tag, pairs[:5]


def embedding_candidates(
    text: str,
    ordered_categories: List[str],
    cat_emb,
    embedder,
    topk: int = 3,
) -> List[str]:
    if not embedder or cat_emb is None or not text:
        return []
    emb = embedder.encode([text], normalize_embeddings=True)
    sims = cosine_sim_matrix(emb, cat_emb)[0]
    pairs = list(zip(ordered_categories, sims))
    pairs.sort(key=lambda x: x[1], reverse=True)
    # only keep those above weak threshold
    out = [c for c, s in pairs[:max(topk, 1)] if s >= EMB_WEAK_THRESHOLD]
    return out[:topk]


# -----------------------
# OpenAI GPT by_id classification (with candidates)
# -----------------------
def _build_openai_client():
    api_key = (os.getenv("OPENAI_API_KEY") or "").strip()
    if not api_key:
        raise RuntimeError(
            "Missing OPENAI_API_KEY. Add it in GitHub repo Settings -> Secrets and variables -> Actions."
        )
    from openai import OpenAI  # type: ignore

    if OPENAI_BASE_URL:
        return OpenAI(api_key=api_key, base_url=OPENAI_BASE_URL)
    return OpenAI(api_key=api_key)


def _openai_chat_complete(client, messages):
    return client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=messages,
        temperature=OPENAI_TEMPERATURE,
        response_format={"type": "json_object"},
    )


def _extract_content(resp) -> str:
    return resp.choices[0].message.content


def _call_with_retries(client, messages, label: str):
    last_err = None
    for attempt in range(1, OPENAI_MAX_RETRIES + 1):
        try:
            return _openai_chat_complete(client, messages)
        except Exception as e:
            last_err = e
            wait = min(10, 2 * attempt)
            print(f"[openai:{label}] error attempt {attempt}: {e} (wait {wait}s)", flush=True)
            time.sleep(wait)
    raise RuntimeError(f"OpenAI request failed after {OPENAI_MAX_RETRIES} attempts: {last_err}")


def gpt_classify_by_id(
    client,
    ordered_categories: List[str],
    cat_desc: Dict[str, str],
    items: List[dict],
    candidates_by_id: Optional[Dict[str, List[str]]] = None,
    tag: str = "gpt_classify",
) -> Tuple[Dict[str, List[str]], Dict[str, float], str]:
    """
    items: [{id,title,text}, ...]
    returns:
      labels_by_id: {id: [cat,...]} (may be missing some ids)
      conf_by_id: {id: float} (optional; missing => not provided)
      raw_response_text
    """
    allowed = ordered_categories
    allowed_set = set(allowed)

    # Make a compact categories block
    cats_block = "\n".join([f"- {c}: {cat_desc.get(c,'')}" for c in allowed])

    # Attach candidates if present
    payload_items = []
    for it in items:
        _id = str(it["id"])
        obj = {
            "id": _id,
            "title": it.get("title", ""),
            "text": it.get("text", ""),
        }
        if candidates_by_id and candidates_by_id.get(_id):
            obj["candidates"] = candidates_by_id[_id]
        payload_items.append(obj)

    system = (
        "You are an expert literature classifier. "
        "Classify each paper into the given categories. Use abstract/text first, title as fallback. "
        "Multi-label is allowed. Do not invent category names."
    )
    user = (
        "分类规则（类别名: 说明）:\n"
        f"{cats_block}\n\n"
        "任务要求（非常重要）:\n"
        "1) 必须只使用给定类别名（不可发明/改写）。\n"
        "2) 每条必须输出，即使不匹配也要输出空数组 []。\n"
        "3) 输出必须是 JSON 且严格为：\n"
        "   {\"by_id\": {\"<id>\": {\"labels\": [\"类A\",\"类B\"], \"confidence\": 0.0}}}\n"
        "   其中 confidence 取 0~1。\n"
        "4) 如果提供 candidates 字段，请优先从 candidates 中选择；若 candidates 都不合适，可返回 []。\n\n"
        f"输入条目:\n{json.dumps(payload_items, ensure_ascii=False)}"
    )

    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    resp = _call_with_retries(client, messages, tag)
    raw = _clean_json_text(_extract_content(resp))

    labels_by_id: Dict[str, List[str]] = {}
    conf_by_id: Dict[str, float] = {}

    try:
        data = json.loads(raw)
    except Exception as e:
        _dump_debug(tag, {"items": payload_items}, raw_text=raw, reason=f"json.loads failed: {e}")
        return labels_by_id, conf_by_id, raw

    by_id = data.get("by_id", {})
    if not isinstance(by_id, dict):
        _dump_debug(tag, {"items": payload_items, "parsed": data}, raw_text=raw, reason="by_id not a dict")
        return labels_by_id, conf_by_id, raw

    for k, v in by_id.items():
        _id = _normalize_cell(k)
        if not _id:
            continue
        if isinstance(v, dict):
            labs = v.get("labels", [])
            conf = v.get("confidence", None)
        else:
            labs = v
            conf = None

        if not isinstance(labs, list):
            labs = []
        clean = []
        for c in labs:
            c = _normalize_cell(c)
            if c in allowed_set:
                clean.append(c)
        labels_by_id[_id] = list(dict.fromkeys(clean))

        if isinstance(conf, (int, float)):
            conf_by_id[_id] = float(conf)

    return labels_by_id, conf_by_id, raw


# -----------------------
# Postprocess: primary label, review
# -----------------------
def choose_primary(labels: List[str], priority: List[str]) -> str:
    s = set(labels or [])
    for p in priority:
        if p in s:
            return p
    return labels[0] if labels else ""


def should_review(route: str, confidence: float, labels: List[str]) -> bool:
    if not labels:
        return True
    if confidence and confidence < REVIEW_CONFIDENCE_THRESHOLD:
        return True
    if route.startswith("gpt") and confidence < 0.50:
        return True
    if "kw:tie" in route:
        return True
    return False


# -----------------------
# DOCX writer (simple, reuse your style)
# -----------------------
def add_hyperlink(paragraph, url: str, text: str, color_hex="1155CC", underline=True):
    if not url:
        paragraph.add_run(text)
        return
    part = paragraph.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    r_pr.append(u)
    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_doc_default_style(doc: Document, font_name="Calibri", font_size_pt=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_divider_line(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("—" * 70)
    run.font.size = Pt(9)


def doi_to_url(doi: str) -> str:
    doi = _normalize_cell(doi)
    if not doi:
        return ""
    if doi.lower().startswith("http"):
        return doi
    return f"https://doi.org/{doi}"


def _write_record_block(doc: Document, row: dict, idx: int):
    title_en = _normalize_cell(row.get("title", ""))
    title_zh = _normalize_cell(row.get("title_zh", ""))
    link = _normalize_cell(row.get("link", ""))
    source = _normalize_cell(row.get("source", ""))
    pub_date = _normalize_cell(row.get("pub_date", "")) or _normalize_cell(row.get("published", ""))
    doi = _normalize_cell(row.get("doi", ""))
    abstract_en = _normalize_cell(row.get("abstract", ""))
    abstract_zh = _normalize_cell(row.get("abstract_zh", ""))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p.add_run(f"[{idx}] ").bold = True
    main_title = title_zh or title_en or "(no title)"
    if link:
        add_hyperlink(p, link, main_title)
    else:
        p.add_run(main_title).bold = True

    if title_zh and title_en and title_zh != title_en:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(f"EN: {title_en}")
        r2.italic = True
        r2.font.size = Pt(10)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(6)
    mr = meta.add_run("Source: ")
    mr.bold = True
    meta.add_run(source or "-")
    meta.add_run("    ")
    mr = meta.add_run("Pub date: ")
    mr.bold = True
    meta.add_run(pub_date or "-")
    if doi:
        meta.add_run("    ")
        mr = meta.add_run("DOI: ")
        mr.bold = True
        add_hyperlink(meta, doi_to_url(doi), doi)

    abs_zh_p = doc.add_paragraph()
    abs_zh_p.paragraph_format.space_before = Pt(0)
    abs_zh_p.paragraph_format.space_after = Pt(4)
    abs_zh_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    abs_zh_p.paragraph_format.line_spacing = 1.15
    abs_zh_p.add_run("摘要（ZH）: ").bold = True
    abs_zh_p.add_run(abstract_zh if abstract_zh else "(empty)")

    abs_en_p = doc.add_paragraph()
    abs_en_p.paragraph_format.space_before = Pt(0)
    abs_en_p.paragraph_format.space_after = Pt(10)
    abs_en_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    abs_en_p.paragraph_format.line_spacing = 1.15
    abs_en_p.add_run("Abstract (EN): ").bold = True
    abs_en_p.add_run(abstract_en if abstract_en else "(empty)")


def df_to_docx_grouped(df: pd.DataFrame, ordered_categories: List[str], output_docx: str, report_title: str):
    doc = Document()
    set_doc_default_style(doc, font_name="Calibri", font_size_pt=11)

    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(report_title)
    r.bold = True
    r.font.size = Pt(18)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    sub_p.add_run(datetime.now().strftime("%Y-%m-%d")).italic = True

    if df.empty:
        doc.add_paragraph("No records.")
        doc.save(output_docx)
        return

    # Build mapping category -> rows
    cat_to_rows: Dict[str, List[dict]] = {c: [] for c in ordered_categories}
    rows = df.to_dict(orient="records")
    for row in rows:
        cats = _normalize_cell(row.get("categories", ""))
        lab = [x.strip() for x in cats.split(";") if x.strip()]
        for c in lab:
            if c in cat_to_rows:
                cat_to_rows[c].append(row)

    appeared = [c for c in ordered_categories if cat_to_rows.get(c)]
    for ci, cat in enumerate(appeared):
        header = doc.add_paragraph()
        header.paragraph_format.space_before = Pt(8)
        header.paragraph_format.space_after = Pt(6)
        hr = header.add_run(f"【{cat}】")
        hr.bold = True
        hr.font.size = Pt(14)

        for j, row in enumerate(cat_to_rows[cat], start=1):
            _write_record_block(doc, row, j)
            if j != len(cat_to_rows[cat]):
                add_divider_line(doc)
        if ci != len(appeared) - 1:
            add_divider_line(doc)

    doc.save(output_docx)


# -----------------------
# XLSX writer
# -----------------------
def write_xlsx(
    df: pd.DataFrame,
    ordered_categories: List[str],
    output_xlsx: str,
    still_missing_ids: List[str],
):
    used = set()
    with pd.ExcelWriter(output_xlsx, engine="openpyxl") as writer:
        # All
        _safe_all = _safe_sheet_name(SHEET_ALL, used)
        df.to_excel(writer, index=False, sheet_name=_safe_all)

        # Unlabeled
        unlabeled = df[df["categories"].str.strip() == ""].copy()
        _safe_un = _safe_sheet_name(SHEET_UNLABELED, used)
        unlabeled.to_excel(writer, index=False, sheet_name=_safe_un)

        # Review
        review = df[df["need_review"].str.lower().isin(["true", "1", "yes", "y", "t"])].copy()
        _safe_rv = _safe_sheet_name(SHEET_REVIEW, used)
        review.to_excel(writer, index=False, sheet_name=_safe_rv)

        # Still missing
        if still_missing_ids:
            miss = df[df["id"].isin(still_missing_ids)].copy()
            _safe_ms = _safe_sheet_name(SHEET_STILL_MISSING, used)
            miss.to_excel(writer, index=False, sheet_name=_safe_ms)

        # Per-category
        for cat in ordered_categories:
            mask = df["categories"].apply(lambda s: cat in [x.strip() for x in str(s).split(";") if x.strip()])
            sub = df[mask].copy()
            if sub.empty:
                continue
            sheet = _safe_sheet_name(cat, used)
            sub.to_excel(writer, index=False, sheet_name=sheet)


# -----------------------
# Main pipeline
# -----------------------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("-i", "--input", default="", help="Input weekly CSV path (default: latest in output/weekly)")
    ap.add_argument("-c", "--classification", default=DEFAULT_CLASSIFICATION_FILE, help="classification.txt path")
    ap.add_argument("--report-title", default="Tech Tracking Digest")
    ap.add_argument("--skip-gpt", action="store_true", help="Skip GPT stage (only keyword+embedding+hard rules)")
    ap.add_argument("--force-gpt", action="store_true", help="Send all items to GPT (ignore auto routes)")
    ap.add_argument("--no-docx", action="store_true", help="Do not generate DOCX")
    args = ap.parse_args()

    # Load categories
    ordered_categories, cat_desc, cat_keywords = parse_classification_file(args.classification)
    print(f"[classify] categories ({len(ordered_categories)}): {ordered_categories}", flush=True)

    # Pick input
    csv_path = Path(args.input) if args.input else pick_latest_weekly_csv(DEFAULT_INPUT_DIR)
    print(f"[io] picked CSV: {csv_path}", flush=True)

    df = pd.read_csv(csv_path, encoding="utf-8-sig", keep_default_na=False)
    df = ensure_base_columns(df)
    print(f"[io] loaded rows: {len(df)}", flush=True)

    # Build stable ids + text_for_classify
    records = df.to_dict(orient="records")
    id_list = []
    text_for_id: Dict[str, str] = {}
    title_for_id: Dict[str, str] = {}
    row_idx_for_id: Dict[str, int] = {}

    for i, row in enumerate(records):
        rid = make_stable_id(row)
        id_list.append(rid)
        row_idx_for_id[rid] = i
        title_for_id[rid] = _normalize_cell(row.get("title", ""))
        text_for_id[rid] = build_text_for_classify(row.get("title", ""), row.get("abstract", ""))

    df["id"] = id_list

    # Prepare embedder and category embeddings
    embedder = try_load_embedder()
    cat_emb = None
    if embedder:
        protos = build_category_prototypes(ordered_categories, cat_desc, cat_keywords)
        try:
            cat_emb = embedder.encode(protos, normalize_embeddings=True)
            print("[embed] category embeddings ready", flush=True)
        except Exception as e:
            print(f"[embed] disabled (encode failed): {e}", flush=True)
            embedder = None
            cat_emb = None

    # Routing results
    labels_by_id: Dict[str, List[str]] = {}
    conf_by_id: Dict[str, float] = {}
    route_by_id: Dict[str, str] = {}
    candidates_by_id: Dict[str, List[str]] = {}

    need_gpt_ids: List[str] = []

    # Stage A/B/C routing
    for rid in id_list:
        text = text_for_id.get(rid, "")
        if not text:
            labels_by_id[rid] = []
            conf_by_id[rid] = 0.0
            route_by_id[rid] = "empty"
            continue

        if args.force_gpt:
            need_gpt_ids.append(rid)
            route_by_id[rid] = "force_gpt"
            continue

        # A) hard overrides first (on empty labels)
        hard_labels, hard_reason = apply_hard_overrides(text, [], ordered_categories)
        if hard_reason and hard_labels:
            labels_by_id[rid] = hard_labels
            conf_by_id[rid] = 0.98
            route_by_id[rid] = hard_reason
            continue

        # B) keyword strong routing
        kw_labels, kw_conf, kw_route = keyword_route(text, ordered_categories, cat_keywords)
        if kw_labels:
            labels_by_id[rid] = kw_labels
            conf_by_id[rid] = kw_conf
            route_by_id[rid] = kw_route
            continue

        # Collect candidates from keyword weak hits
        kw_cands = keyword_candidates(text, ordered_categories, cat_keywords, topk=EMB_TOPK_CANDIDATES)

        # C) embedding strong routing / candidates
        emb_labels, emb_conf, emb_route, _top = embedding_route_single(text, ordered_categories, cat_emb, embedder)
        if emb_labels:
            labels_by_id[rid] = emb_labels
            conf_by_id[rid] = emb_conf
            route_by_id[rid] = emb_route
            continue

        emb_cands = embedding_candidates(text, ordered_categories, cat_emb, embedder, topk=EMB_TOPK_CANDIDATES)

        # Merge candidates (kw + emb) de-duped
        merged = []
        for x in (kw_cands + emb_cands):
            if x and x not in merged:
                merged.append(x)
        if merged:
            candidates_by_id[rid] = merged[:EMB_TOPK_CANDIDATES]
            route_by_id[rid] = f"candidates({','.join(candidates_by_id[rid])})"
        else:
            route_by_id[rid] = "no_candidates"

        need_gpt_ids.append(rid)

    print(f"[route] need_gpt={len(need_gpt_ids)} / total={len(id_list)}", flush=True)

    # GPT stage
    still_missing_ids: List[str] = []
    if need_gpt_ids and not args.skip_gpt:
        client = _build_openai_client()

        start = time.time()
        done = 0
        total = len(need_gpt_ids)

        # First pass GPT in batches
        for batch_ids in _chunked(need_gpt_ids, CLASSIFY_BATCH_SIZE):
            items = [{"id": rid, "title": title_for_id.get(rid, ""), "text": text_for_id.get(rid, "")} for rid in batch_ids]
            cands = {rid: candidates_by_id.get(rid, []) for rid in batch_ids if candidates_by_id.get(rid)}

            out_labels, out_conf, raw = gpt_classify_by_id(
                client,
                ordered_categories=ordered_categories,
                cat_desc=cat_desc,
                items=items,
                candidates_by_id=cands if cands else None,
                tag="gpt_classify",
            )

            # Merge results
            for rid in batch_ids:
                if rid in out_labels:
                    labels_by_id[rid] = out_labels[rid]
                    conf_by_id[rid] = float(out_conf.get(rid, 0.55))  # default mid if not provided
                    route_by_id[rid] = route_by_id.get(rid, "gpt") + "|gpt"
                else:
                    # missing -> keep for retry
                    pass

            done += len(batch_ids)
            elapsed = time.time() - start
            rate = done / elapsed if elapsed > 0 else 0.0
            eta = (total - done) / rate if rate > 0 else 0.0
            print(f"[gpt#1] {done}/{total} elapsed={_fmt_secs(elapsed)} rate={rate:.2f} it/s ETA={_fmt_secs(eta)}", flush=True)

        missing_after_first = [rid for rid in need_gpt_ids if rid not in labels_by_id]
        if missing_after_first:
            print(f"[gpt#1] missing ids: {len(missing_after_first)}", flush=True)

        # Retry missing once
        if missing_after_first and GPT_RETRY_MISSING_ONCE:
            retry_ids = missing_after_first
            print(f"[gpt#2] retry missing (once), n={len(retry_ids)}", flush=True)

            for batch_ids in _chunked(retry_ids, max(1, min(CLASSIFY_BATCH_SIZE, 8))):
                items = [{"id": rid, "title": title_for_id.get(rid, ""), "text": text_for_id.get(rid, "")} for rid in batch_ids]
                cands = {rid: candidates_by_id.get(rid, []) for rid in batch_ids if candidates_by_id.get(rid)}

                out_labels, out_conf, raw = gpt_classify_by_id(
                    client,
                    ordered_categories=ordered_categories,
                    cat_desc=cat_desc,
                    items=items,
                    candidates_by_id=cands if cands else None,
                    tag="gpt_retry_missing",
                )

                for rid in batch_ids:
                    if rid in out_labels:
                        labels_by_id[rid] = out_labels[rid]
                        conf_by_id[rid] = float(out_conf.get(rid, 0.55))
                        route_by_id[rid] = route_by_id.get(rid, "gpt") + "|gpt_retry"
                    else:
                        # still missing after retry -> record
                        pass

            still_missing_ids = [rid for rid in retry_ids if rid not in labels_by_id]
            if still_missing_ids:
                print(f"[gpt#2] still missing after retry: {len(still_missing_ids)}", flush=True)

    elif need_gpt_ids and args.skip_gpt:
        print("[gpt] skipped by --skip-gpt; unresolved items will remain unlabeled", flush=True)
        still_missing_ids = []

    # Final postprocess: hard overrides + normalize to string join
    priority = ordered_categories[:]  # use file order as priority unless you want custom
    # If you have a category called "其他", make it lowest priority by default
    if "其他" in priority:
        priority = [c for c in priority if c != "其他"] + ["其他"]

    final_cats_col = []
    final_primary = []
    final_conf = []
    final_route = []
    final_review = []

    for rid in id_list:
        text = text_for_id.get(rid, "")
        labs = labels_by_id.get(rid, [])
        conf = float(conf_by_id.get(rid, 0.0))
        route = route_by_id.get(rid, "")

        # Apply hard overrides again on the chosen labels
        labs2, reason = apply_hard_overrides(text, labs, ordered_categories)
        if reason:
            # override confidence & route
            labs = labs2
            conf = max(conf, 0.90)
            route = (route + "|" if route else "") + reason

        # Ensure labels are valid
        allowed = set(ordered_categories)
        labs = [x for x in labs if x in allowed]
        labs = list(dict.fromkeys(labs))

        primary = choose_primary(labs, priority)

        # If empty labels and you have "其他", optionally assign it (comment out if you want truly empty)
        if not labs and "其他" in ordered_categories:
            labs = ["其他"]
            primary = "其他"
            conf = max(conf, 0.40)
            route = (route + "|" if route else "") + "fallback:其他"

        need_review = should_review(route, conf, labs)

        final_cats_col.append(";".join(labs))
        final_primary.append(primary)
        final_conf.append(f"{conf:.3f}")
        final_route.append(route)
        final_review.append("True" if need_review else "False")

    df["categories"] = final_cats_col
    df["primary_category"] = final_primary
    df["confidence"] = final_conf
    df["route"] = final_route
    df["need_review"] = final_review

    # Output naming (match your style: *_translated.xlsx / *.docx)
    out_xlsx = csv_path.with_name(csv_path.stem + "_translated.xlsx")
    out_docx = out_xlsx.with_suffix(".docx")

    write_xlsx(df, ordered_categories, str(out_xlsx), still_missing_ids)
    print(f"[io] wrote XLSX: {out_xlsx}", flush=True)

    if not args.no_docx:
        df_to_docx_grouped(df, ordered_categories, str(out_docx), report_title=args.report_title)
        print(f"[io] wrote DOCX: {out_docx}", flush=True)
    else:
        print("[io] skipped DOCX by --no-docx", flush=True)


if __name__ == "__main__":
    main()
