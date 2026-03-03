# -*- coding: utf-8 -*-
"""
test_classify_only.py (by_id + retry-missing)

Only classify (no translation) + output xlsx/docx.

Key upgrades:
1) Batch classification uses stable IDs and returns a mapping: {"by_id": {"<id>": ["类A","类B"], ...}}
2) If any id is missing from model output, we collect them and ask GPT again (single retry) to fill gaps.
3) If still missing after retry, those records go to a separate Excel sheet: "仍缺失标签"
4) Always write something for every id: empty list means "no label" (未匹配)
5) Debug artifacts for failures are dumped to output/debug (items, raw, reason)
"""

import argparse
import json
import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

print("[boot] test_classify_only.py started", flush=True)

# ============================
# OpenAI configuration
# ============================
OPENAI_MODEL = (os.getenv("OPENAI_MODEL") or "").strip() or "gpt-4o-mini"
OPENAI_TEMPERATURE = float(os.getenv("OPENAI_TEMPERATURE", "0"))
OPENAI_MAX_RETRIES = int(os.getenv("OPENAI_MAX_RETRIES", "3"))
OPENAI_BASE_URL = (os.getenv("OPENAI_BASE_URL") or "").strip()

CLASSIFY_BATCH_SIZE = int(os.getenv("CLASSIFY_BATCH_SIZE", "12"))
CLASSIFY_TEXT_MAX_CHARS = int(os.getenv("CLASSIFY_TEXT_MAX_CHARS", "1200"))
DEBUG_DUMP_DIR = os.getenv("DEBUG_DUMP_DIR", "output/debug")

DEFAULT_CLASSIFICATION_FILE = "classification.txt"

# Excel sheet for records still missing after 2nd call
SHEET_STILL_MISSING = "仍缺失标签"


def _fmt_secs(s: float) -> str:
    s = max(0, int(s))
    h = s // 3600
    m = (s % 3600) // 60
    sec = s % 60
    if h > 0:
        return f"{h}h {m}m {sec}s"
    if m > 0:
        return f"{m}m {sec}s"
    return f"{sec}s"


def _chunked(items, n):
    for i in range(0, len(items), n):
        yield items[i : i + n]


def _normalize_cell(v):
    if v is None:
        return ""
    return str(v).strip()


def _clean_json_text(s: str) -> str:
    s = (s or "").strip()
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.IGNORECASE)
        s = re.sub(r"\s*```$", "", s)
    return s.strip()


# ============================
# OpenAI client
# ============================
def _build_openai_client():
    api_key = (os.getenv("OPENAI_API_KEY") or "").strip()
    if not api_key:
        raise RuntimeError(
            "Missing OPENAI_API_KEY in environment. "
            "For GitHub Actions, add it in repo Settings -> Secrets and variables -> Actions."
        )

    from openai import OpenAI

    if OPENAI_BASE_URL:
        return OpenAI(api_key=api_key, base_url=OPENAI_BASE_URL)
    return OpenAI(api_key=api_key)


def _openai_chat_complete(client, messages):
    if not OPENAI_MODEL:
        raise RuntimeError("OPENAI_MODEL is empty. Set OPENAI_MODEL or rely on default.")
    return client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=messages,
        temperature=OPENAI_TEMPERATURE,
        response_format={"type": "json_object"},
    )


def _extract_content(resp):
    return resp.choices[0].message.content


def _call_with_retries(client, messages, label: str):
    last_err = None
    for attempt in range(1, OPENAI_MAX_RETRIES + 1):
        try:
            return _openai_chat_complete(client, messages)
        except Exception as e:
            last_err = e
            wait = min(10, 2 * attempt)
            print(f"[openai:{label}] error on attempt {attempt}: {e} (wait {wait}s)", flush=True)
            time.sleep(wait)
    raise RuntimeError(f"OpenAI request failed after {OPENAI_MAX_RETRIES} attempts: {last_err}")


# ============================
# Input discovery / normalization
# ============================
def _extract_date_from_name(filename: str):
    stem = Path(filename).stem
    m = re.search(r"(\d{4})[-_]?(\d{2})[-_]?(\d{2})", stem)
    if not m:
        return None
    y, mo, d = map(int, m.groups())
    try:
        return datetime(y, mo, d).date()
    except ValueError:
        return None


def pick_latest_weekly_csv(folder="output/weekly") -> Path:
    folder = Path(folder)
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder.resolve()}")

    prefixes = ["weekly_news_with_abstract_", "news_with_abstract_"]
    candidates = []
    for p in folder.glob("*.csv"):
        if "_translated" in p.stem:
            continue
        if any(p.name.startswith(pref) for pref in prefixes):
            candidates.append(p)

    if not candidates:
        raise FileNotFoundError(f"No weekly CSV found in {folder.resolve()}")

    with_dates = []
    without_dates = []
    for p in candidates:
        d = _extract_date_from_name(p.name)
        if d is not None:
            with_dates.append((d, p))
        else:
            without_dates.append(p)

    if with_dates:
        with_dates.sort(key=lambda x: (x[0], x[1].stat().st_mtime), reverse=True)
        return with_dates[0][1]
    return max(without_dates, key=lambda p: p.stat().st_mtime)


def ensure_base_columns(df: pd.DataFrame) -> pd.DataFrame:
    needed = [
        "title",
        "link",
        "published",
        "source",
        "pub_date",
        "doi",
        "abstract",
        "abstract_source",
        "must_have_abstract",
        "title_zh",
        "abstract_zh",
        "categories",
    ]
    for c in needed:
        if c not in df.columns:
            df[c] = ""
        df[c] = df[c].fillna("").astype(str)
    return df


# ============================
# Classification rules
# ============================
def load_classification_rules(path: str):
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"classification file not found: {p.resolve()}")

    rules = []
    for raw in p.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if "：" in line:
            name, desc = line.split("：", 1)
        elif ":" in line:
            name, desc = line.split(":", 1)
        else:
            name, desc = line, ""
        name = name.strip()
        desc = desc.strip()
        if name:
            rules.append((name, desc))

    if not rules:
        raise RuntimeError(f"No valid category lines parsed from {p.resolve()}")
    return rules


def _dump_debug(batch_items, content, reason: str, tag: str = "classify"):
    outdir = Path(DEBUG_DUMP_DIR)
    outdir.mkdir(parents=True, exist_ok=True)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    stem = f"{tag}_fail_{ts}"

    (outdir / f"{stem}_reason.txt").write_text(reason, encoding="utf-8")
    (outdir / f"{stem}_items.json").write_text(
        json.dumps(batch_items, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (outdir / f"{stem}_raw.txt").write_text(content or "", encoding="utf-8")
    print(f"[debug] dumped failure artifacts to {outdir.resolve()}", flush=True)


def _coerce_labels_list(x) -> List[str]:
    if not isinstance(x, list):
        return []
    out = []
    for v in x:
        s = _normalize_cell(v)
        if s:
            out.append(s)
    # de-dup preserve order
    return list(dict.fromkeys(out))


def classify_batch_by_id(
    client,
    rules: List[Tuple[str, str]],
    items: List[Dict],
    label: str = "classify_by_id",
) -> Tuple[Dict[str, List[str]], str]:
    """
    Returns (by_id_map, raw_content).

    by_id_map keys are str(ids); values are list of category names (possibly empty list).
    If a given id is missing from output, it simply won't be present in by_id_map.
    """
    allowed = [name for name, _ in rules]
    allowed_set = set(allowed)
    categories_block = "\n".join([f"- {name}: {desc}" for name, desc in rules])

    system = (
        "You are an expert literature classifier. "
        "Classify each paper using abstract first, title as fallback. "
        "Multi-label is allowed. Do not invent category names."
    )

    # Important: make schema as explicit as possible.
    user = (
        "分类规则（类别名: 说明）:\n"
        f"{categories_block}\n\n"
        "任务要求:\n"
        "1) 每个条目可属于多个类别。\n"
        "2) 必须只使用给定类别名（不可发明/改写）。\n"
        "3) 若都不匹配，必须返回空数组 []。\n"
        "4) 必须为每个输入条目的 id 返回一个键（不允许漏掉）。\n"
        "5) 仅返回JSON，且严格符合：\n"
        " {\"by_id\": {\"<id>\": [\"类别A\", \"类别B\"], \"<id2>\": []}}\n\n"
        f"输入条目（每项含 id/title/text）：\n{json.dumps(items, ensure_ascii=False)}"
    )

    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    resp = _call_with_retries(client, messages, label)
    raw = _clean_json_text(_extract_content(resp))

    try:
        data = json.loads(raw)
    except Exception as e:
        _dump_debug(items, raw, f"json.loads failed: {e}", tag=label)
        return {}, raw

    by_id = data.get("by_id", {})
    if not isinstance(by_id, dict):
        _dump_debug(items, raw, f"by_id not a dict: {type(by_id)}", tag=label)
        return {}, raw

    normalized: Dict[str, List[str]] = {}
    for k, v in by_id.items():
        k_str = _normalize_cell(k)
        if not k_str:
            continue
        labs = _coerce_labels_list(v)
        # filter invalid labels
        labs = [x for x in labs if x in allowed_set]
        normalized[k_str] = labs

    return normalized, raw


def classify_records(df: pd.DataFrame, rules: List[Tuple[str, str]]):
    """
    Returns:
      df_out: df with 'categories' column
      labels_list: List[List[str]] aligned to df rows
      still_missing_row_indices: List[int] (rows missing after 2nd call)
    """
    df = ensure_base_columns(df.copy())

    # build items with stable IDs
    items_all = []
    row_ids = [] # str id per row (stable)
    for idx, row in df.iterrows():
        abstract = _normalize_cell(row.get("abstract", ""))
        title = _normalize_cell(row.get("title", ""))
        text = abstract if abstract else title
        text = _normalize_cell(text)
        if len(text) > CLASSIFY_TEXT_MAX_CHARS:
            text = text[:CLASSIFY_TEXT_MAX_CHARS]
        # stable id: row index + 1 (or you can use hash; this is deterministic per run)
        rid = str(int(idx) + 1)
        row_ids.append(rid)
        items_all.append({"id": rid, "title": title, "text": text})

    client = _build_openai_client()
    allowed_set = set([name for name, _ in rules])

    # First pass classification (batched)
    by_id_total: Dict[str, List[str]] = {}
    missing_ids_first: List[str] = []

    start = time.time()
    total = len(items_all)
    processed = 0

    for batch in _chunked(items_all, CLASSIFY_BATCH_SIZE):
        # Optional: handle empty text deterministically to reduce model weirdness
        non_empty = [it for it in batch if _normalize_cell(it.get("text", "")) != ""]
        empty = [it for it in batch if it not in non_empty]

        by_id_map, _raw = classify_batch_by_id(client, rules, non_empty, label="classify_by_id")
        # fill in empty-text ones as no-label
        for it in empty:
            by_id_map[str(it["id"])] = []

        # merge
        for _id, labs in by_id_map.items():
            if not isinstance(labs, list):
                labs = []
            labs = [x for x in labs if x in allowed_set]
            by_id_total[_id] = list(dict.fromkeys(labs))

        # track missing within this batch (only for non-empty items; empty already filled)
        batch_ids = [str(it["id"]) for it in batch]
        for bid in batch_ids:
            if bid not in by_id_total:
                # shouldn't happen if model obeys; but we'll handle
                missing_ids_first.append(bid)

        processed += len(batch)
        elapsed = time.time() - start
        rate = processed / elapsed if elapsed > 0 else 0.0
        eta = (total - processed) / rate if rate > 0 else 0.0
        print(
            f"[classify#1] {processed}/{total} | elapsed={_fmt_secs(elapsed)} | rate={rate:.2f} items/s | ETA={_fmt_secs(eta)}",
            flush=True,
        )

    # Determine missing after first pass
    missing_after_first = [rid for rid in row_ids if rid not in by_id_total]
    missing_after_first = list(dict.fromkeys(missing_after_first))
    if missing_after_first:
        print(f"[classify#1] missing ids after first pass: {len(missing_after_first)}", flush=True)

    # Second pass: ask GPT only for missing ids
    missing_after_second: List[str] = []
    if missing_after_first:
        retry_items = []
        id_to_item = {str(it["id"]): it for it in items_all}
        for mid in missing_after_first:
            it = id_to_item.get(mid)
            if it:
                retry_items.append(it)

        # Do a single retry in batches (same size or smaller)
        for batch in _chunked(retry_items, max(1, min(CLASSIFY_BATCH_SIZE, 8))):
            by_id_map, _raw = classify_batch_by_id(client, rules, batch, label="classify_by_id_retry_missing")
            for _id, labs in by_id_map.items():
                labs = [x for x in labs if x in allowed_set]
                by_id_total[_id] = list(dict.fromkeys(labs))

        # compute still missing
        missing_after_second = [mid for mid in missing_after_first if mid not in by_id_total]
        if missing_after_second:
            print(f"[classify#2] still missing after second pass: {len(missing_after_second)}", flush=True)

    # Build labels_list aligned to df rows; if truly missing (after 2nd), set [].
    labels_list: List[List[str]] = []
    for rid in row_ids:
        labs = by_id_total.get(rid, [])
        if not isinstance(labs, list):
            labs = []
        labels_list.append(labs)

    df["categories"] = [";".join(x) for x in labels_list]

    # Convert missing ids to row indices (0-based df position)
    # rid = str(idx+1) so idx = int(rid)-1
    still_missing_row_indices = []
    for mid in missing_after_second:
        try:
            still_missing_row_indices.append(int(mid) - 1)
        except Exception:
            continue

    return df, labels_list, still_missing_row_indices


# ============================
# DOCX helpers (same layout as before)
# ============================
def add_hyperlink(paragraph, url: str, text: str, color_hex="1155CC", underline=True):
    if not url:
        paragraph.add_run(text)
        return

    part = paragraph.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    r_pr.append(u)
    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_doc_default_style(doc: Document, font_name="Calibri", font_size_pt=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_divider_line(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("—" * 70)
    run.font.size = Pt(9)


def doi_to_url(doi: str) -> str:
    doi = _normalize_cell(doi)
    if not doi:
        return ""
    if doi.lower().startswith("http"):
        return doi
    return f"https://doi.org/{doi}"


def strip_leading_abstract(text: str) -> str:
    if not text:
        return text
    s = text.lstrip()
    return re.sub(r"^(abstract)\s*[:.\-–—]*\s*", "", s, flags=re.IGNORECASE)


def abstract_to_runs(abstract: str):
    if not abstract:
        return [("(empty)", False)]

    abstract = strip_leading_abstract(abstract)
    text = abstract.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")

    chunks = []
    pending_space = False
    for line in lines:
        s = line.strip()
        if s == "":
            pending_space = True
            continue
        if pending_space and chunks:
            chunks.append((" ", False))
            pending_space = False

        if s.isdigit():
            chunks.append((s, True))
        else:
            if chunks:
                prev_text, prev_sub = chunks[-1]
                if not prev_sub and prev_text and not prev_text.endswith(" ") and prev_text[-1] not in {"-", "−", "/"}:
                    chunks.append((" ", False))
            chunks.append((s, False))

    cleaned = []
    for t, sub in chunks:
        if cleaned and t == " " and cleaned[-1][0] == " ":
            continue
        cleaned.append((t, sub))
    return cleaned


def add_abstract_with_subscripts(paragraph, abstract: str):
    for t, is_sub in abstract_to_runs(abstract):
        run = paragraph.add_run(t)
        if is_sub:
            run.font.subscript = True


def _write_record_block(doc, row, idx):
    title_en = _normalize_cell(row.get("title", ""))
    title_zh = _normalize_cell(row.get("title_zh", "")) # may be empty in classify-only mode
    link = _normalize_cell(row.get("link", ""))
    source = _normalize_cell(row.get("source", ""))
    pub_date = _normalize_cell(row.get("pub_date", "")) or _normalize_cell(row.get("published", ""))
    doi = _normalize_cell(row.get("doi", ""))
    abstract_en = _normalize_cell(row.get("abstract", ""))
    abstract_zh = _normalize_cell(row.get("abstract_zh", "")) # may be empty
    abstract_source = _normalize_cell(row.get("abstract_source", ""))
    must_have_abstract = _normalize_cell(row.get("must_have_abstract", ""))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p.add_run(f"[{idx}] ").bold = True
    main_title = title_zh or title_en or "(no title)"
    if link:
        add_hyperlink(p, link, main_title)
    else:
        p.add_run(main_title).bold = True

    if title_en and title_zh and title_en != title_zh:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(f"EN: {title_en}")
        r2.italic = True
        r2.font.size = Pt(10)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(6)
    mr = meta.add_run("Source: ")
    mr.bold = True
    meta.add_run(source or "-")
    meta.add_run(" ")
    mr = meta.add_run("Pub date: ")
    mr.bold = True
    meta.add_run(pub_date or "-")
    if doi:
        meta.add_run(" ")
        mr = meta.add_run("DOI: ")
        mr.bold = True
        add_hyperlink(meta, doi_to_url(doi), doi)

    extra_bits = []
    if abstract_source:
        extra_bits.append(f"abstract_source={abstract_source}")
    if _normalize_cell(must_have_abstract).lower() in {"true", "1", "yes", "y", "t"}:
        extra_bits.append("must_have_abstract=True")
    if extra_bits:
        extra = doc.add_paragraph()
        extra.paragraph_format.space_before = Pt(0)
        extra.paragraph_format.space_after = Pt(6)
        er = extra.add_run("Notes: ")
        er.bold = True
        extra.add_run("; ".join(extra_bits))

    abs_zh_p = doc.add_paragraph()
    abs_zh_p.paragraph_format.space_before = Pt(0)
    abs_zh_p.paragraph_format.space_after = Pt(4)
    abs_zh_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    abs_zh_p.paragraph_format.line_spacing = 1.15
    abs_zh_p.add_run("摘要（ZH）: ").bold = True
    add_abstract_with_subscripts(abs_zh_p, abstract_zh if abstract_zh else "(empty)")

    abs_en_p = doc.add_paragraph()
    abs_en_p.paragraph_format.space_before = Pt(0)
    abs_en_p.paragraph_format.space_after = Pt(10)
    abs_en_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    abs_en_p.paragraph_format.line_spacing = 1.15
    abs_en_p.add_run("Abstract (EN): ").bold = True
    add_abstract_with_subscripts(abs_en_p, abstract_en if abstract_en else "(empty)")


def df_to_word_bilingual_grouped(df: pd.DataFrame, labels, ordered_categories, output_docx: str, report_title="Tech Tracking Digest (Classify-Only Test)"):
    doc = Document()
    set_doc_default_style(doc, font_name="Calibri", font_size_pt=11)

    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(report_title)
    r.bold = True
    r.font.size = Pt(18)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    sub_p.add_run(datetime.now().strftime("%Y-%m-%d")).italic = True

    if df.empty:
        doc.add_paragraph("No records found in CSV.")
        doc.save(output_docx)
        return

    records = df.to_dict(orient="records")
    cat_to_indices = {c: [] for c in ordered_categories}
    for i, cats in enumerate(labels):
        for c in cats:
            if c in cat_to_indices:
                cat_to_indices[c].append(i)

    appeared_categories = [c for c in ordered_categories if cat_to_indices.get(c)]
    for ci, cat in enumerate(appeared_categories):
        header = doc.add_paragraph()
        header.paragraph_format.space_before = Pt(8)
        header.paragraph_format.space_after = Pt(6)
        hr = header.add_run(f"【{cat}】")
        hr.bold = True
        hr.font.size = Pt(14)

        indices = cat_to_indices[cat]
        for j, ridx in enumerate(indices, start=1):
            _write_record_block(doc, records[ridx], j)
            if j != len(indices):
                add_divider_line(doc)
        if ci != len(appeared_categories) - 1:
            add_divider_line(doc)

    doc.save(output_docx)


# ============================
# XLSX output
# ============================
def _safe_sheet_name(name: str, used: set) -> str:
    base = re.sub(r"[\\/*?:]", "_", _normalize_cell(name)) or "Sheet"
    base = base[:31]
    candidate = base
    i = 2
    while candidate in used:
        suffix = f"_{i}"
        candidate = base[: 31 - len(suffix)] + suffix
        i += 1
    used.add(candidate)
    return candidate


def write_grouped_xlsx(
    df: pd.DataFrame,
    labels: List[List[str]],
    ordered_categories: List[str],
    output_xlsx: str,
    still_missing_row_indices: List[int],
):
    if df.empty:
        df.to_excel(output_xlsx, index=False, sheet_name="Empty")
        return

    cat_to_indices = {c: [] for c in ordered_categories}
    for i, cats in enumerate(labels):
        for c in cats:
            if c in cat_to_indices:
                cat_to_indices[c].append(i)

    used = set()
    with pd.ExcelWriter(output_xlsx, engine="openpyxl") as writer:
        # per-category sheets
        for cat in ordered_categories:
            idxs = cat_to_indices.get(cat, [])
            if not idxs:
                continue
            sheet = _safe_sheet_name(cat, used)
            sub = df.iloc[idxs].copy()
            sub.to_excel(writer, index=False, sheet_name=sheet)

        # sheet for still-missing after second pass
        if still_missing_row_indices:
            sheet = _safe_sheet_name(SHEET_STILL_MISSING, used)
            sub = df.iloc[still_missing_row_indices].copy()
            sub.to_excel(writer, index=False, sheet_name=sheet)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("-i", "--input", default="", help="Input weekly CSV path (default: latest in output/weekly)")
    parser.add_argument("-c", "--classification", default=DEFAULT_CLASSIFICATION_FILE, help="Classification rules file")
    parser.add_argument("--report-title", default="Tech Tracking Digest (Classify-Only Test)")
    args = parser.parse_args()

    if args.input:
        csv_path = Path(args.input)
    else:
        csv_path = pick_latest_weekly_csv(folder="output/weekly")

    print(f"[io] Picked CSV: {csv_path}", flush=True)
    df = pd.read_csv(csv_path, encoding="utf-8-sig", keep_default_na=False)
    df = ensure_base_columns(df)
    print(f"[io] Loaded rows: {len(df)}", flush=True)

    rules = load_classification_rules(args.classification)
    ordered_categories = [name for name, _ in rules]
    print(f"[classify] loaded categories: {ordered_categories}", flush=True)
    print(
        f"[classify] model={OPENAI_MODEL} batch={CLASSIFY_BATCH_SIZE} text_max_chars={CLASSIFY_TEXT_MAX_CHARS}",
        flush=True,
    )

    df_out, labels, still_missing = classify_records(df, rules)

    output_xlsx = csv_path.with_name(csv_path.stem + "_classify_only.xlsx")
    write_grouped_xlsx(df_out, labels, ordered_categories, str(output_xlsx), still_missing)
    print(f"[io] Wrote XLSX: {output_xlsx}", flush=True)

    output_docx = output_xlsx.with_suffix(".docx")
    df_to_word_bilingual_grouped(df_out, labels, ordered_categories, str(output_docx), report_title=args.report_title)
    print(f"[io] Wrote DOCX: {output_docx}", flush=True)

    if still_missing:
        print(f"[warn] still missing labels after 2nd pass: {len(still_missing)} (see sheet '{SHEET_STILL_MISSING}')", flush=True)


if __name__ == "__main__":
    main()
