# -*- coding: utf-8 -*-
"""
Weekly CSV -> translate (Marian) -> translated CSV -> bilingual DOCX

Features:
- Progress + ETA (titles/abstracts separately)
- Abstract truncation before translation (CPU-friendly)
- Only translate missing fields (title_zh/abstract_zh empty)
- Output:
    output/weekly/<base>_translated.csv
    output/weekly/<base>_translated.docx

Model:
  Helsinki-NLP/opus-mt-en-zh

Notes:
- No caching of translation results beyond writing into translated CSV.
- Model caching is handled by HF cache directory (YAML cache step recommended).
"""

import csv
import re
import time
from pathlib import Path
from datetime import datetime

import pandas as pd

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline


# ============================
# Translation configuration
# ============================
MODEL_NAME = "Helsinki-NLP/opus-mt-en-zh"

# ✅ 关键：控制翻译规模（CPU 上跑得动）
MAX_ABSTRACT_CHARS_TO_TRANSLATE = 800   # 建议 600~1200；越小越快
BATCH_SIZE_TITLE = 16                   # 标题短，batch 可大
BATCH_SIZE_ABSTRACT = 2                 # 摘要长，batch 小更稳
MAX_LENGTH_TITLE = 128
MAX_LENGTH_ABSTRACT = 256               # 512 会更慢；周报用 256 通常够


# ============================
# Progress helpers
# ============================
def _fmt_secs(s: float) -> str:
    s = max(0, int(s))
    h = s // 3600
    m = (s % 3600) // 60
    sec = s % 60
    if h > 0:
        return f"{h}h {m}m {sec}s"
    if m > 0:
        return f"{m}m {sec}s"
    return f"{sec}s"


def translate_with_progress(translator, texts, batch_size, max_length, label: str):
    """
    Translate list[str] with progress + ETA.
    """
    total = len(texts)
    if total == 0:
        print(f"[translate:{label}] nothing to translate")
        return []

    print(f"[translate:{label}] total={total}, batch_size={batch_size}, max_length={max_length}")
    out = []
    start = time.time()
    last_print = start

    for i in range(0, total, batch_size):
        batch = texts[i:i + batch_size]

        # The actual translation call
        res = translator(batch, max_length=max_length)
        out.extend([x["translation_text"] for x in res])

        done = min(i + batch_size, total)
        now = time.time()

        # Print every batch (so you always see progress in GitHub logs)
        elapsed = now - start
        rate = done / elapsed if elapsed > 0 else 0.0
        eta = (total - done) / rate if rate > 0 else 0.0

        print(f"[translate:{label}] {done}/{total} | elapsed={_fmt_secs(elapsed)} | "
              f"rate={rate:.2f} items/s | ETA={_fmt_secs(eta)}")

        last_print = now

    return out


# ============================
# Translator builder
# ============================
def build_translator():
    print(f"[model] loading tokenizer+model: {MODEL_NAME}")
    tok = AutoTokenizer.from_pretrained(MODEL_NAME)
    model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME)
    print("[model] building pipeline (CPU)")
    return pipeline(
        "translation",
        model=model,
        tokenizer=tok,
        device=-1,  # CPU
        # Optional: silence the FutureWarning by setting it explicitly:
        clean_up_tokenization_spaces=True
    )


# ============================
# Data translation
# ============================
def enrich_translation(df: pd.DataFrame) -> pd.DataFrame:
    """
    Add/Fill title_zh and abstract_zh columns using Marian.
    - Translate title if title_zh empty.
    - Translate abstract if abstract exists and abstract_zh empty.
    """
    for col in ["title", "abstract"]:
        if col not in df.columns:
            df[col] = ""
        df[col] = df[col].fillna("").astype(str)

    if "title_zh" not in df.columns:
        df["title_zh"] = ""
    if "abstract_zh" not in df.columns:
        df["abstract_zh"] = ""

    df["title_zh"] = df["title_zh"].fillna("").astype(str)
    df["abstract_zh"] = df["abstract_zh"].fillna("").astype(str)

    # Determine what needs translation
    need_title_idx = df.index[
        (df["title"].str.strip() != "") & (df["title_zh"].str.strip() == "")
    ].tolist()

    need_abs_idx = df.index[
        (df["abstract"].str.strip() != "") & (df["abstract_zh"].str.strip() == "")
    ].tolist()

    print(f"[plan] titles to translate: {len(need_title_idx)}")
    print(f"[plan] abstracts to translate: {len(need_abs_idx)} "
          f"(abstract truncation={MAX_ABSTRACT_CHARS_TO_TRANSLATE} chars)")

    if not need_title_idx and not need_abs_idx:
        print("[plan] nothing to translate, skip model load")
        return df

    translator = build_translator()

    # ---- Titles ----
    if need_title_idx:
        titles = df.loc[need_title_idx, "title"].tolist()
        zh_titles = translate_with_progress(
            translator,
            titles,
            batch_size=BATCH_SIZE_TITLE,
            max_length=MAX_LENGTH_TITLE,
            label="title"
        )
        df.loc[need_title_idx, "title_zh"] = zh_titles

    # ---- Abstracts ----
    if need_abs_idx:
        abstracts = df.loc[need_abs_idx, "abstract"].tolist()

        # ✅ truncate BEFORE translation to control runtime
        truncated = []
        for a in abstracts:
            a = a or ""
            a = a.strip()
            if len(a) > MAX_ABSTRACT_CHARS_TO_TRANSLATE:
                a = a[:MAX_ABSTRACT_CHARS_TO_TRANSLATE]
            truncated.append(a)

        zh_abs = translate_with_progress(
            translator,
            truncated,
            batch_size=BATCH_SIZE_ABSTRACT,
            max_length=MAX_LENGTH_ABSTRACT,
            label="abstract"
        )
        df.loc[need_abs_idx, "abstract_zh"] = zh_abs

    return df


# ============================
# Word helpers (same as yours + bilingual)
# ============================
def add_hyperlink(paragraph, url: str, text: str, color_hex="1155CC", underline=True):
    if not url:
        paragraph.add_run(text)
        return

    part = paragraph.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex)
    r_pr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    r_pr.append(u)

    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_doc_default_style(doc: Document, font_name="Calibri", font_size_pt=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_divider_line(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("—" * 70)
    run.font.size = Pt(9)


def safe_get(row: dict, key: str) -> str:
    return (row.get(key) or "").strip()


def doi_to_url(doi: str) -> str:
    doi = (doi or "").strip()
    if not doi:
        return ""
    if doi.lower().startswith("http"):
        return doi
    return f"https://doi.org/{doi}"


def normalize_datetime_str(s: str) -> str:
    return (s or "").strip()


def is_truthy_flag(s: str) -> bool:
    s = (s or "").strip().lower()
    return s in {"true", "1", "yes", "y", "t"}


def strip_leading_abstract(text: str) -> str:
    if not text:
        return text
    s = text.lstrip()
    s = re.sub(r"^(abstract)\s*[:.\-–—]*\s*", "", s, flags=re.IGNORECASE)
    return s


def abstract_to_runs(abstract: str):
    if not abstract:
        return [("(empty)", False)]

    abstract = strip_leading_abstract(abstract)

    text = abstract.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")

    chunks = []
    pending_space = False

    for line in lines:
        s = line.strip()
        if s == "":
            pending_space = True
            continue

        if pending_space and chunks:
            chunks.append((" ", False))
            pending_space = False

        if s.isdigit():
            chunks.append((s, True))
        else:
            if chunks:
                prev_text, prev_sub = chunks[-1]
                if not prev_sub and prev_text and not prev_text.endswith(" "):
                    if prev_text[-1] not in {"-", "−", "/"}:
                        chunks.append((" ", False))
            chunks.append((s, False))

    cleaned = []
    for t, sub in chunks:
        if cleaned and t == " " and cleaned[-1][0] == " ":
            continue
        cleaned.append((t, sub))
    return cleaned


def add_abstract_with_subscripts(paragraph, abstract: str):
    for t, is_sub in abstract_to_runs(abstract):
        run = paragraph.add_run(t)
        if is_sub:
            run.font.subscript = True


# ============================
# CSV picker
# ============================
def _extract_date_from_name(filename: str):
    stem = Path(filename).stem
    m = re.search(r"(\d{4})[-_]?(\d{2})[-_]?(\d{2})", stem)
    if not m:
        return None
    y, mo, d = map(int, m.groups())
    try:
        return datetime(y, mo, d).date()
    except ValueError:
        return None


def pick_latest_weekly_csv(folder="output/weekly") -> Path:
    """
    Support prefixes:
      - weekly_news_with_abstract_
      - news_with_abstract_
    Excludes *_translated.csv
    """
    folder = Path(folder)
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder.resolve()}")

    prefixes = ["weekly_news_with_abstract_", "news_with_abstract_"]
    candidates = []
    for p in folder.glob("*.csv"):
        if p.name.endswith("_translated.csv"):
            continue
        if any(p.name.startswith(pref) for pref in prefixes):
            candidates.append(p)

    if not candidates:
        raise FileNotFoundError(f"No weekly CSV found in {folder.resolve()}")

    with_dates = []
    without_dates = []
    for p in candidates:
        d = _extract_date_from_name(p.name)
        if d is not None:
            with_dates.append((d, p))
        else:
            without_dates.append(p)

    if with_dates:
        with_dates.sort(key=lambda x: (x[0], x[1].stat().st_mtime), reverse=True)
        return with_dates[0][1]

    return max(without_dates, key=lambda p: p.stat().st_mtime)


# ============================
# DOCX generation (bilingual)
# ============================
def df_to_word_bilingual(df: pd.DataFrame, output_docx: str, report_title: str = "Tech Tracking Digest"):
    doc = Document()
    set_doc_default_style(doc, font_name="Calibri", font_size_pt=11)

    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(report_title)
    r.bold = True
    r.font.size = Pt(18)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    sub_p.add_run(datetime.now().strftime("%Y-%m-%d")).italic = True

    if df.empty:
        doc.add_paragraph("No records found in CSV.")
        doc.save(output_docx)
        return

    needed_cols = [
        "title", "title_zh", "link", "source", "pub_date", "published", "doi",
        "abstract", "abstract_zh", "abstract_source", "must_have_abstract"
    ]
    for c in needed_cols:
        if c not in df.columns:
            df[c] = ""

    records = df.to_dict(orient="records")

    for idx, row in enumerate(records, start=1):
        title_en = safe_get(row, "title")
        title_zh = safe_get(row, "title_zh")
        link = safe_get(row, "link")
        source = safe_get(row, "source")
        pub_date = safe_get(row, "pub_date") or safe_get(row, "published")
        doi = safe_get(row, "doi")
        abstract_en = safe_get(row, "abstract")
        abstract_zh = safe_get(row, "abstract_zh")
        abstract_source = safe_get(row, "abstract_source")
        must_have_abstract = safe_get(row, "must_have_abstract")

        # Header
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        p.add_run(f"[{idx}] ").bold = True
        main_title = title_zh or title_en or "(no title)"
        if link:
            add_hyperlink(p, link, main_title)
        else:
            p.add_run(main_title).bold = True

        # EN title line if both exist
        if title_en and title_zh and title_en.strip() != title_zh.strip():
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(4)
            p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r2 = p2.add_run(f"EN: {title_en}")
            r2.italic = True
            r2.font.size = Pt(10)

        # Meta
        meta = doc.add_paragraph()
        meta.paragraph_format.space_before = Pt(0)
        meta.paragraph_format.space_after = Pt(6)
        meta.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        mr = meta.add_run("Source: ")
        mr.bold = True
        meta.add_run(source or "-")

        meta.add_run("    ")
        mr = meta.add_run("Pub date: ")
        mr.bold = True
        meta.add_run(normalize_datetime_str(pub_date) or "-")

        if doi:
            meta.add_run("    ")
            mr = meta.add_run("DOI: ")
            mr.bold = True
            add_hyperlink(meta, doi_to_url(doi), doi)

        extra_bits = []
        if abstract_source:
            extra_bits.append(f"abstract_source={abstract_source}")
        if is_truthy_flag(must_have_abstract):
            extra_bits.append("must_have_abstract=True")

        if extra_bits:
            extra = doc.add_paragraph()
            extra.paragraph_format.space_before = Pt(0)
            extra.paragraph_format.space_after = Pt(6)
            er = extra.add_run("Notes: ")
            er.bold = True
            extra.add_run("; ".join(extra_bits))

        # Abstract bilingual
        abs_zh_p = doc.add_paragraph()
        abs_zh_p.paragraph_format.space_before = Pt(0)
        abs_zh_p.paragraph_format.space_after = Pt(4)
        abs_zh_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        abs_zh_p.paragraph_format.line_spacing = 1.15
        abs_zh_p.add_run("摘要（ZH）: ").bold = True
        add_abstract_with_subscripts(abs_zh_p, abstract_zh if abstract_zh else "(empty)")

        abs_en_p = doc.add_paragraph()
        abs_en_p.paragraph_format.space_before = Pt(0)
        abs_en_p.paragraph_format.space_after = Pt(10)
        abs_en_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        abs_en_p.paragraph_format.line_spacing = 1.15
        abs_en_p.add_run("Abstract (EN): ").bold = True
        add_abstract_with_subscripts(abs_en_p, abstract_en if abstract_en else "(empty)")

        if idx != len(records):
            add_divider_line(doc)

    doc.save(output_docx)


# ============================
# Main
# ============================
if __name__ == "__main__":
    weekly_dir = Path("output/weekly")
    csv_path = pick_latest_weekly_csv(folder=str(weekly_dir))
    print(f"[io] Picked CSV: {csv_path}")

    df = pd.read_csv(csv_path, encoding="utf-8-sig", keep_default_na=False)

    # Translate with progress + ETA
    df = enrich_translation(df)

    # Write translated CSV
    translated_csv_path = csv_path.with_name(csv_path.stem + "_translated.csv")
    df.to_csv(translated_csv_path, index=False, encoding="utf-8-sig")
    print(f"[io] Wrote translated CSV: {translated_csv_path}")

    # Write DOCX
    docx_path = translated_csv_path.with_suffix(".docx")
    TITLE = "Tech Tracking Digest"
    df_to_word_bilingual(df, str(docx_path), report_title=TITLE)
    print(f"[io] Wrote DOCX: {docx_path}")
